import os
import pydicom

from docx import Document                       
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE       
from docx.shared import Pt, Inches, RGBColor  


class SummaryReport:
    def __init__(self):
        self.p_info = None
        self.plane_info = ["Axial", "Coronal", "Sagittal"]
        self.img_paths = None
        self.result_info = None
        
    def set_personal_info(self, dcm_path: str):
        _dcm_file = pydicom.dcmread(dcm_path, force=True)
        self.p_info = {
            0: ["환자 ID", "이름", "성별", "나이", "생년월일"],
            1: [
                _dcm_file.PatientID, 
                _dcm_file.PatientName, 
                _dcm_file.PatientSex, 
                _dcm_file.PatientAge, 
                _dcm_file.PatientBirthDate
            ]
        }
        
    # gradcam path
    def set_image_paths(self, result_path: str):
        from glob import glob 
        self.img_paths = glob(f"{result_path}/*.png")
        
    # percentage
    def set_result_info(self, result: list[int]):
        self.result_info = {
            0: ["Abnormal (비정상)", f"{result[0]}%", ""],
            1: ["ACL tear (전방십자인대)", f"{result[1]}%", ""],
            2: ["Meniscus tear (반달연골)", f"{result[2]}%", ""]
        }

    def export_to_docx(self):
        document = Document("base_document.docx")

        # personal info title
        p_info_paragraph = document.add_paragraph("환자 정보")
        p_info_paragraph.paragraph_format.space_after = Pt(5)
        run = p_info_paragraph.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0, 0, 0)

        # personal info table
        p_info_table = document.add_table(rows=2, cols=5)
        p_info_table.style = "custom_style"
        p_info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.tables[-1]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):       
                # set text and alignment
                for para in cell.paragraphs:
                    if i == 0:
                        para.add_run(self.p_info[i][j]).bold = True 
                    else:
                        para.add_run(self.p_info[i][j])
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # set font            
                    for run in para.runs:
                        run.font.name = "맑은 고딕"
                        run.font.size = Pt(12)

        document.add_paragraph()

        # result title
        result_paragraph = document.add_paragraph("검사 결과")
        result_paragraph.paragraph_format.space_after = Pt(5)
        run = result_paragraph.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0, 0, 0)

        # plane table
        plane_table = document.add_table(rows=1, cols=3)
        plane_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        plane_table.style = "custom_style2"

        table = document.tables[-1]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):       
                # set text and alignment
                for para in cell.paragraphs:
                    para.add_run(self.plane_info[i][j]).bold = True 
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # set font            
                    for run in para.runs:
                        run.font.name = "맑은 고딕"
                        run.font.size = Pt(12)
                        
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.space_before = Pt(5)
                cell.paragraphs[0].paragraph_format.space_after = Pt(5)

        # image table
        img_paragraph = document.add_paragraph()
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_paragraph.paragraph_format.space_after = Pt(10)
        for img_path in self.img_paths:
            run = img_paragraph.add_run()
            run.add_picture(img_path, width=Inches(2.5))
            img_paragraph.add_run().add_text(" "*3)

        explain_paragraph = document.add_paragraph()
        explain_paragraph.add_run("""무릎 MRI를 활용한 질병 예측 모델에서는 GradCAM을 통해 각 plane별로 중요한 부위를 시각화하고, CAM Score를 계산하여 모델이 판단한 가장 중요한 이미지를 도출합니다. 이를 통해 모델의 예측 결과를 시각적으로""")
        explain_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        explain_paragraph.add_run(" 해석하고, CAM Score를 통해 예측의 신뢰도를 확인할 수 있습니다.")
        explain_paragraph.paragraph_format.space_after = Pt(10)

        # diseases table
        result_table = document.add_table(rows=3, cols=3)
        result_table.style = "custom_style"
        result_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.tables[-1]
        first_col = table.columns[0]
        first_col.width = Inches(8)

        # find idxs that have score over 50
        idxs = [[i, float(k[1][:-1])] for i, k in self.result_info.items()]
        idxs = [i[0] for i in list(filter(lambda x: x[1] >= 50, idxs))]
        
        for i, row in enumerate(table.rows):
            res = self.result_info[i]
            if i in idxs:
                res[2] = "✔"
            for j, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    para.add_run(res[j])
                    if j == 0: 
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    for run in para.runs:
                        run.font.name = "맑은 고딕"
                        run.font.size = Pt(12)
                        
                if i in idxs:
                    cell.paragraphs[0].runs[0].bold = True
                    
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.space_before = Pt(3)
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)
                
        caption_para = document.add_paragraph(style="Caption")
        caption_para.add_run(" * ✔ 표시의 경우 모델이 예측한 결과 해당 질병이 있을 확률이 높다는 것을 의미합니다.")
        caption_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        save_docs_path = f"{self.p_info[1][0]}_auto_report.docx"
        if os.path.exists(save_docs_path):
            os.remove(save_docs_path)    
        document.save(save_docs_path)
        
    
    
if __name__ == '__main__':
    summary_report = SummaryReport()
    summary_report.set_personal_info("output_dicom.dcm")
    summary_report.set_image_paths([
        "../backend/sampleimages/axial_0000.png", 
        "../backend/sampleimages/axial_0001.png", 
        "../backend/sampleimages/axial_0002.png"
    ])
    summary_report.set_result_info({
        0: ["Abnormal (비정상)", "50%", ""],
        1: ["ACL tear (전방십자인대)", "20%", ""],
        2: ["Meniscus tear (반달연골)", "60%", ""]
    })
    summary_report.export_to_docx()
    